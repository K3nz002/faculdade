from docx import Document
from docx.shared import Inches
import requests
from bs4 import BeautifulSoup
from io import BytesIO

# doc = Document()

# doc.add_heading('Estes são os seus dados', 0)

# nome = input("Qual o seu nome? ")
# p = doc.add_paragraph("Seu nome é ")
# p.add_run(nome).bold = True
# p.add_run('.')

# doc.add_heading('A seguir outros detalhes:', level=1)

# idade = input("Qual a sua idade? ")
# doc.add_paragraph('Você tem {} anos de idade'.format(idade), style='List Bullet')

# telefone = input("Qual é o seu telefone? ")
# doc.add_paragraph('Seu telefone é o {}'.format(telefone), style='List Bullet')

# endereco = input("Qual é o seu endereço? ")
# doc.add_paragraph('Seu endereço é {}'.format(endereco), style='List Bullet')

# foto = input("Digite o caminho (path) completo de um arquivo com sua foto? ")
# doc.add_picture(foto, width=Inches(6), height=Inches(4))

# # Salvar arquivo no diretório atual
# doc.save('meudoc.docx')

''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''''

numero = int(input('Tabuada de qual número? '))
doc = Document()

doc.add_heading('Tabuada de {}'.format(numero), 0)

tab = doc.add_table(rows=1, cols=5)
tab.style="Medium List 1 Accent 2"
cels = tab.rows[0].cells
cels[0].text = 'Núm'
cels[1].text = "Opr"
cels[2].text = "Rep"
cels[3].text = "=" 
cels[4].text = "Res"

for rep in range(11):
   dados = tab.add_row().cells
   dados[0].text = str(numero)
   dados[1].text = 'x'
   dados[2].text = str(rep)
   dados[3].text = '='
   dados[4].text = str(numero * rep)

doc.save('tabuada{}.docx'.format(numero))

